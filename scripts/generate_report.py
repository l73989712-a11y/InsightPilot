from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "《Python Web应用实训》实训报告-InsightPilot.docx"
IMG = ROOT / "docs" / "images"

TITLE = "基于 Flask 与大语言模型的智能数据分析与可视化平台设计与实现"


def set_run_font(run, east="宋体", ascii_font="Times New Roman", size=12, bold=None, color=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = paragraph.add_run("- ")
    set_run_font(r1, size=10.5)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    r2 = paragraph.add_run(" -")
    set_run_font(r2, size=10.5)


def restart_page_numbering(section, start=1):
    sectPr = section._sectPr
    pg_num = sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sectPr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)


def add_body(doc, text, bold_prefix=None, center=False, no_indent=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    if not no_indent and not center:
        pf.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size)
    return p


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            set_run_font(r, size=11.5)
        if not p.runs:
            r = p.add_run(item)
            set_run_font(r, size=11.5)
        else:
            p.runs[0].text = item
    return


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F5F8")
    set_cell_margins(cell, 100, 120, 100, 120)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(code.strip())
    set_run_font(r, east="等线", ascii_font="Consolas", size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_picture(doc, relpath, width_cm, caption=None):
    path = IMG / relpath
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        add_caption(doc, caption)


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, east="黑体", ascii_font="Arial", size=16, bold=True)
    elif level == 2:
        set_run_font(r, east="黑体", ascii_font="Arial", size=14, bold=True)
    else:
        set_run_font(r, east="黑体", ascii_font="Arial", size=12, bold=True)
    return p


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], "DCE6F1")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def cover(doc):
    sec = doc.sections[0]
    set_page(sec)
    sec.top_margin = Cm(2.0)
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《Python Web 应用实训》")
    set_run_font(r, east="黑体", ascii_font="Arial", size=22, bold=True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实 训 报 告")
    set_run_font(r, east="黑体", ascii_font="Arial", size=26, bold=True)
    p.paragraph_format.space_after = Pt(60)

    info = [
        ("学    院：", "信息科学与工程学院"),
        ("专    业：", "人工智能"),
        ("班    级：", "2024 级本科 X 班（请填写）"),
        ("姓    名：", "____________________"),
        ("学    号：", "____________________"),
        ("组员姓名：", "____________________（单人完成可删除）"),
        ("组员学号：", "____________________"),
        ("项目名称：", TITLE),
        ("实训日期：", "2026 年 5 月 15 日至 2026 年 5 月 30 日"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.4)
    table.columns[1].width = Cm(12.5)
    for i, (a,b) in enumerate(info):
        table.rows[i].height = Cm(1.05)
        set_cell_text(table.cell(i,0), a, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i,1), b, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        for cell in table.rows[i].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for border_name in ["top","left","bottom","right","insideH","insideV"]:
                borders = tcPr.first_child_found_in("w:tcBorders")
                if borders is None:
                    borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
                border = borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}"); borders.append(border)
                border.set(qn("w:val"), "nil")
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("齐鲁师范学院信息科学与工程学院")
    set_run_font(r, east="黑体", ascii_font="Arial", size=14, bold=True)
    doc.add_page_break()


def record_sheet(doc, member_no):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("齐鲁师范学院信息科学与工程学院"); set_run_font(r,east="黑体",size=11,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Python Web 应用实训记录单"); set_run_font(r,east="黑体",size=18,bold=True)
    p.paragraph_format.space_after=Pt(18)
    rows=[
        ["学号", "________________", "姓名", "________________", "性别", "____"],
        ["项目名称", TITLE, "", "", "", ""],
        ["项目描述", "本项目面向 CSV/Excel 数据分析场景，实现用户登录、数据上传、质量检测、数据清洗、统计分析、可视化、自然语言分析和报告生成。大模型仅负责生成结构化分析计划与解释，真实数据计算由 Pandas 和 Scikit-learn 完成。", "", "", "", ""],
        ["分工描述", "第 %d 位成员分工（请按实际情况填写）：需求分析、后端开发、数据分析、前端页面、测试或报告编写。" % member_no, "", "", "", ""],
        ["成绩", "□优秀  □良好  □中等  □合格  □不合格", "", "", "", ""],
    ]
    table=doc.add_table(rows=5,cols=6); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.autofit=False
    widths=[2.2,3.2,1.7,3.0,1.7,2.0]
    for row in table.rows:
        for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    for i in range(6):
        set_cell_text(table.cell(0,i),rows[0][i],bold=i%2==0,size=10.5,align=WD_ALIGN_PARAGRAPH.CENTER)
    # merged rows
    for r_idx in [1,2,3,4]:
        table.cell(r_idx,1).merge(table.cell(r_idx,5))
        set_cell_text(table.cell(r_idx,0),rows[r_idx][0],bold=True,size=10.5,align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(r_idx,1),rows[r_idx][1],size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT)
    table.rows[1].height=Cm(1.6); table.rows[2].height=Cm(6.0); table.rows[3].height=Cm(5.0); table.rows[4].height=Cm(2.0)
    doc.add_page_break()


def abstract_page(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("摘  要"); set_run_font(r,east="黑体",size=18,bold=True)
    p.paragraph_format.space_after=Pt(18)
    texts=[
        "随着数据规模不断增长，普通用户在面对 CSV 或 Excel 数据时，往往难以快速完成质量检查、清洗、统计计算和图表解释。传统管理类 Python Web 实训项目虽然能够展示增删改查能力，但与 Python 数据分析课程的衔接不够紧密，也难以体现当前大语言模型技术的发展趋势。为此，本项目设计并实现了 InsightPilot——基于 Flask 与大语言模型的智能数据分析与可视化平台。",
        "系统采用 Flask 构建 Web 应用，使用 MySQL 5.7 保存用户、数据集元信息、清洗记录、分析任务、AI 会话和报告；使用 Pandas、NumPy 与 Scikit-learn 完成真实的数据载入、质量检测、缺失值处理、异常值处理、分组聚合、相关分析、线性回归和 K-Means 聚类；前端使用 Bootstrap 与本地 Plotly.js 完成交互界面和图表展示。系统同时提供规则模式和兼容 OpenAI 接口的大模型模式。大模型负责将自然语言问题转换为结构化 JSON 分析计划，并根据真实计算结果生成文字解释，后端通过动作白名单、字段校验和参数校验阻止任意代码执行。",
        "项目实现了完整的数据分析闭环：用户上传数据后，系统自动生成健康度报告；用户可执行多种清洗操作并保留记录；统计工作台支持常见分析任务；AI 助手允许用户使用自然语言提出分析需求；最终可生成包含数据概况、分析记录和 AI 总结的 HTML 报告。测试结果表明，核心流程能够正确完成，项目规模适合作为一至两人的 Python Web 实训项目，并具有较好的展示性和扩展性。",
    ]
    for t in texts: add_body(doc,t)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
    r=p.add_run("关键词："); set_run_font(r,east="黑体",size=12,bold=True)
    r=p.add_run("Flask；数据分析；Pandas；数据可视化；大语言模型；MySQL"); set_run_font(r,size=12)
    doc.add_page_break()


def toc_pages(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("目  录"); set_run_font(r,east="黑体",size=18,bold=True)
    p.paragraph_format.space_after=Pt(16)
    entries=[
        ("第一章 概述",0,"C1"),
        ("1.1 项目背景",1,"H11"),("1.2 项目意义",1,"H12"),("1.3 项目主要工作",1,"H13"),("1.4 报告结构",1,"H14"),
        ("第二章 需求分析",0,"C2"),("2.1 用户与业务需求",1,"H21"),("2.2 功能需求",1,"H22"),("2.3 非功能需求",1,"H23"),("2.4 用例分析",1,"H24"),("2.5 可行性分析",1,"H25"),
        ("第三章 相关技术",0,"C3"),("3.1 Flask Web 框架",1,"H31"),("3.2 MySQL 与 SQLAlchemy",1,"H32"),("3.3 Pandas 与 NumPy",1,"H33"),("3.4 Scikit-learn",1,"H34"),("3.5 Bootstrap 与 Plotly.js",1,"H35"),("3.6 大语言模型与安全调用",1,"H36"),
        ("第四章 概要设计",0,"C4"),("4.1 系统总体架构",1,"H41"),("4.2 功能模块设计",1,"H42"),("4.3 AI 分析流程",1,"H43"),("4.4 数据库设计",1,"H44"),("4.5 部署结构",1,"H45"),
        ("第五章 详细设计与实现",0,"C5"),("5.1 工程结构与应用工厂",1,"H51"),("5.2 用户认证模块",1,"H52"),("5.3 数据集管理模块",1,"H53"),("5.4 数据质量检测模块",1,"H54"),("5.5 数据清洗模块",1,"H55"),("5.6 统计分析与可视化模块",1,"H56"),("5.7 AI 数据助手模块",1,"H57"),("5.8 报告生成模块",1,"H58"),("5.9 用户界面实现",1,"H59"),
        ("第六章 系统测试",0,"C6"),("6.1 测试环境与方法",1,"H61"),("6.2 功能测试",1,"H62"),("6.3 自动化测试结果",1,"H63"),("6.4 数据分析结果验证",1,"H64"),("6.5 安全与异常测试",1,"H65"),
        ("第七章 总结与展望",0,"C7"),("参考文献",0,"REF"),("附录",0,"APP"),
    ]
    table=doc.add_table(rows=0,cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for text,level,key in entries:
        cells=table.add_row().cells; cells[0].width=Cm(14.3); cells[1].width=Cm(1.4)
        p0=cells[0].paragraphs[0]; p0.paragraph_format.left_indent=Cm(0.7*level); p0.paragraph_format.space_after=Pt(1)
        r0=p0.add_run(text); set_run_font(r0,east="黑体" if level==0 else "宋体",size=11,bold=(level==0))
        p1=cells[1].paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p1.paragraph_format.space_after=Pt(1)
        r1=p1.add_run(f"[[PAGE:{key}]]"); set_run_font(r1,size=11)
        for cell in cells:
            tcPr=cell._tc.get_or_add_tcPr(); borders=OxmlElement("w:tcBorders")
            for name in ("top","left","bottom","right","insideH","insideV"):
                e=OxmlElement(f"w:{name}"); e.set(qn("w:val"),"nil"); borders.append(e)
            tcPr.append(borders)
    doc.add_page_break()


def configure_body_section(doc):
    section=doc.add_section(WD_SECTION.NEW_PAGE); set_page(section)
    section.header.is_linked_to_previous=False; section.footer.is_linked_to_previous=False
    hp=section.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=hp.add_run("齐鲁师范学院信息科学与工程学院"); set_run_font(r,east="黑体",size=10.5,bold=True)
    pPr=hp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"6"); bottom.set(qn("w:space"),"4"); bottom.set(qn("w:color"),"000000"); pBdr.append(bottom); pPr.append(pBdr)
    fp=section.footer.paragraphs[0]; add_page_field(fp); restart_page_numbering(section,1)


def build():
    doc=Document()
    styles=doc.styles
    styles["Normal"].font.name="Times New Roman"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); styles["Normal"].font.size=Pt(12)
    for lvl in [1,2,3]:
        style=styles[f"Heading {lvl}"]
        style.font.name="Arial"; style._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体")
        style.font.bold=True
        style.font.size=Pt({1:16,2:14,3:12}[lvl])

    cover(doc); record_sheet(doc,1); record_sheet(doc,2); abstract_page(doc); toc_pages(doc); configure_body_section(doc)

    # Chapter 1
    add_heading(doc,"第一章 概述",1)
    add_heading(doc,"1.1 项目背景",2)
    for t in [
        "Python 数据分析课程主要训练数据读取、清洗、统计、可视化和简单机器学习能力，而 Python Web 应用实训要求将 Python 程序包装成可供用户操作的网络应用。若只实现传统学生管理、图书管理或库存管理系统，虽然能够覆盖登录和数据库增删改查，但无法充分体现数据分析课程的知识延伸。因此，本项目选择将 Web 工程能力与数据分析能力结合，并引入当前具有代表性的大语言模型交互方式。",
        "在实际学习和工作中，数据通常以 CSV 或 Excel 文件形式存在。使用者首先需要理解字段含义，之后才能判断缺失值、重复值和异常值，再选择适当的统计指标和图表。对于缺少编程经验的用户，上述过程存在较高门槛。大语言模型能够理解自然语言，但其生成的统计数字可能不可靠，因此系统不能让模型直接替代计算程序，而应将模型定位为“意图解析器和结果解释器”。",
        "InsightPilot 的核心思想是把自然语言请求转换为受控的结构化分析计划。例如用户输入“按专业统计平均成绩并绘制柱状图”，模型只返回 group_aggregate 动作、分组字段、数值字段和聚合方法。后端确认字段真实存在且动作属于白名单后，再调用 Pandas 执行计算。这样既保留了大模型交互的便利性，又保证结果来源可追溯。",
    ]:
        add_body(doc,t)
    add_heading(doc,"1.2 项目意义",2)
    for t in [
        "从课程学习角度看，本项目覆盖文件操作、函数与类、Pandas DataFrame、数据清洗、分组聚合、相关分析、可视化、回归和聚类等知识，同时增加 Flask 路由、模板、会话、数据库 ORM、文件上传和 Web 安全等内容，能够形成较完整的综合实践。",
        "从应用价值看，系统为非专业用户提供统一的数据分析入口。用户不需要编写 Python 代码，即可查看数据健康度、执行常见清洗、生成交互图表和保存分析历史。规则模式无需联网即可演示；配置兼容接口后，又可以体验真实大模型对自然语言的理解与总结能力。",
        "从工程价值看，系统采用模块化蓝图、应用工厂、服务层和 ORM 模型，避免将全部代码集中在单一文件中。项目还包含自动化测试、环境检查脚本、示例数据、安装手册和答辩说明，便于后续维护与扩展。",
    ]: add_body(doc,t)
    add_heading(doc,"1.3 项目主要工作",2)
    add_bullets(doc,[
        "完成用户注册、登录、退出和数据权限隔离；",
        "完成 CSV、XLS、XLSX 文件上传、编码识别、字段推断和数据预览；",
        "实现缺失值、重复值、IQR 异常值检测及数据健康度评分；",
        "实现去重、删除缺失、均值/中位数/众数填补、文本清理、异常值删除和标准化；",
        "实现描述统计、频数统计、分组聚合、相关分析、时间趋势、透视表、线性回归和 K-Means；",
        "实现 Plotly.js 交互图表、分析历史、自然语言分析与 HTML 报告；",
        "实现动作白名单、字段校验、CSRF 防护、密码摘要和文件类型限制。",
    ])
    add_heading(doc,"1.4 报告结构",2)
    add_body(doc,"本报告共七章。第一章说明项目背景和意义；第二章分析用户、功能和非功能需求；第三章介绍主要技术；第四章描述系统架构、模块、数据流程和数据库；第五章说明各模块的详细实现；第六章给出测试过程和数据结果；第七章总结项目并提出后续改进方向。")

    # Chapter 2
    add_heading(doc,"第二章 需求分析",1,page_break=True)
    add_heading(doc,"2.1 用户与业务需求",2)
    for t in [
        "系统主要面向需要分析中小型表格数据的学生和普通用户。用户希望在浏览器中完成数据上传、质量检查、清洗、统计和图表生成，而不必搭建 Jupyter Notebook 或编写脚本。用户还希望系统能够记住已完成的分析任务，并将多个结果汇总为报告。",
        "项目规模按照一至两人的课程实训控制。单个文件默认不超过 20 MB，数据行数建议不超过 5 万行。系统不追求分布式计算和海量数据处理，而是强调功能完整、逻辑清晰、可以演示和能够回答教师提问。",
    ]: add_body(doc,t)
    add_heading(doc,"2.2 功能需求",2)
    req_rows=[
        ("F01","用户注册登录","用户名唯一，密码摘要存储，登录后访问业务页面","高"),
        ("F02","数据集上传","支持 CSV、XLS、XLSX，限制扩展名与文件大小","高"),
        ("F03","数据预览","显示字段、类型、前若干行和数据规模","高"),
        ("F04","质量检测","统计缺失、重复、唯一值、异常值和健康度","高"),
        ("F05","数据清洗","去重、缺失处理、文本清理、异常处理、标准化","高"),
        ("F06","统计分析","描述、频数、聚合、相关、回归、聚类等","高"),
        ("F07","数据可视化","柱状图、折线图、散点图和热力图","高"),
        ("F08","AI 分析","自然语言转 JSON 计划，执行后生成解释","高"),
        ("F09","历史记录","保存清洗、分析和 AI 会话","中"),
        ("F10","报告生成","生成并导出 HTML 数据分析报告","中"),
    ]
    add_table(doc,["编号","需求名称","说明","优先级"],req_rows,[1.4,2.8,9.2,1.5],9.5)
    add_heading(doc,"2.3 非功能需求",2)
    non_rows=[
        ("易用性","主要操作通过表单和按钮完成；错误信息使用中文提示。"),
        ("安全性","密码使用哈希摘要；启用 CSRF；数据按用户隔离；AI 动作白名单。"),
        ("可靠性","异常时回滚数据库事务；文件和参数校验；提供离线规则模式。"),
        ("性能","5 万行以内的常见分析应在可接受时间完成；预览只读取部分行。"),
        ("兼容性","推荐 Python 3.10/3.11、MySQL 5.7、Chrome/Edge。"),
        ("可维护性","蓝图拆分路由，服务层封装算法，配置通过环境变量管理。"),
        ("可测试性","分析服务可独立测试，集成测试覆盖完整业务流程。"),
    ]
    add_table(doc,["属性","要求"],non_rows,[3.2,12.5],10)
    add_heading(doc,"2.4 用例分析",2)
    add_body(doc,"普通用户是系统的主要参与者，可以完成注册登录、数据集管理、质量检测、数据清洗、统计分析、AI 提问和报告生成。管理员功能属于扩展内容，可在后续增加用户数量、磁盘占用和任务统计。")
    add_picture(doc,"diagrams/use_case.png",8.4,"图 2-1 系统用例图")
    add_heading(doc,"2.5 可行性分析",2)
    add_body(doc,"技术可行性方面，Flask、Pandas、SQLAlchemy 和 Scikit-learn 均具有成熟的 Python 接口，项目不需要训练大模型，也不需要 GPU。经济可行性方面，开发工具和主要依赖均可免费使用；在规则模式下不产生接口费用。操作可行性方面，系统采用浏览器界面，演示流程清晰。时间可行性方面，功能被划分为基础必做和扩展部分，适合在实训周期内分阶段完成。")

    # Chapter 3
    add_heading(doc,"第三章 相关技术",1,page_break=True)
    add_heading(doc,"3.1 Flask Web 框架",2)
    add_body(doc,"Flask 是轻量级 Python Web 框架，适合中小型课程项目。系统使用应用工厂 create_app 创建 Flask 实例，通过蓝图拆分 auth、datasets、analysis、ai 和 reports 路由。模板采用 Jinja2，用户会话由 Flask-Login 管理，表单和异步请求使用 Flask-WTF 提供的 CSRF 保护。")
    add_code(doc,"""
from flask import Flask
from .extensions import db, login_manager, csrf

def create_app(config_class=Config):
    app = Flask(__name__)
    app.config.from_object(config_class)
    db.init_app(app)
    login_manager.init_app(app)
    csrf.init_app(app)
    app.register_blueprint(auth_bp)
    app.register_blueprint(datasets_bp, url_prefix='/datasets')
    return app
""")
    add_heading(doc,"3.2 MySQL 与 SQLAlchemy",2)
    add_body(doc,"MySQL 5.7 用于保存结构化业务数据。系统没有把原始表格文件以二进制形式写入数据库，而是保存路径、文件大小、行列数、字段信息和质量报告。SQLAlchemy 将数据表映射为 Python 模型，减少手写 SQL 并便于维护。字段结构、分析参数和结果使用 JSON 字符串存入 LONGTEXT，从而避免依赖较新的数据库特性。")
    add_heading(doc,"3.3 Pandas 与 NumPy",2)
    add_body(doc,"Pandas 是数据处理核心。read_csv 和 read_excel 负责读取文件，DataFrame 提供缺失值统计、重复检测、分组聚合、透视表和相关系数计算。NumPy 用于数值类型判断、空值转换和结果序列化。系统对 CSV 依次尝试 UTF-8-SIG、UTF-8、GBK 和 GB18030 编码，以提高中文数据兼容性。")
    add_heading(doc,"3.4 Scikit-learn",2)
    add_body(doc,"Scikit-learn 用于线性回归和 K-Means 聚类。线性回归返回系数、截距和决定系数 R²；聚类返回每个样本的簇标签、聚类中心和惯性指标。项目只选择教材中易于解释的算法，避免把系统扩展为复杂机器学习平台。")
    add_heading(doc,"3.5 Bootstrap 与 Plotly.js",2)
    add_body(doc,"Bootstrap 提供响应式栅格、按钮、表单和卡片组件；项目把 Bootstrap 和 Plotly.js 文件放入 static/vendor 目录，减少答辩现场对外网的依赖。后端返回统一图表配置，前端 JavaScript 将柱状图、折线图、散点图和热力图转换为 Plotly 图表。")
    add_heading(doc,"3.6 大语言模型与安全调用",2)
    add_body(doc,"系统支持规则引擎和兼容 OpenAI Chat Completions 的接口。规则模式通过关键词和字段名称匹配生成分析计划，适合离线演示；接口模式向模型提供字段结构和严格提示词，要求只返回 JSON。无论使用哪种模式，后端都不会执行模型生成的 Python 代码。")
    add_code(doc,"""
ALLOWED_ACTIONS = {
    'describe', 'value_counts', 'group_aggregate',
    'correlation', 'time_trend', 'pivot_table',
    'linear_regression', 'kmeans'
}

if plan['action'] not in ALLOWED_ACTIONS:
    raise ValueError('模型返回了不安全的分析动作')
""")
    add_body(doc,"除了动作白名单，系统还验证字段是否存在、聚合方法是否合法、聚类数量是否在范围内、两个回归字段是否不同。该方案把大模型输出视为不可信输入，只有通过校验的计划才能进入真实计算层。")

    # Chapter 4
    add_heading(doc,"第四章 概要设计",1,page_break=True)
    add_heading(doc,"4.1 系统总体架构",2)
    add_body(doc,"系统采用浏览器、Web 应用、业务服务和存储四层结构。浏览器负责表单交互和图表显示；Flask 负责路由、模板、认证和接口；服务层完成文件处理、数据分析、AI 计划和报告生成；MySQL 与文件系统分别保存业务记录和原始数据。")
    add_picture(doc,"diagrams/system_architecture.png",16.0,"图 4-1 系统总体架构")
    add_heading(doc,"4.2 功能模块设计",2)
    add_body(doc,"系统功能按照业务边界拆分为用户管理、数据集管理、质量检测、数据清洗、统计分析、可视化、AI 助手和报告中心。模块之间通过数据集编号关联，避免页面直接操作任意文件路径。")
    add_picture(doc,"diagrams/function_modules.png",16.0,"图 4-2 功能模块结构")
    add_heading(doc,"4.3 AI 分析流程",2)
    add_body(doc,"AI 分析不是一次简单的聊天请求，而是由字段结构提取、计划生成、计划校验、真实计算、图表生成、结果解释和记录保存组成。计划校验失败时立即终止流程，并向用户返回中文错误信息。")
    add_picture(doc,"diagrams/ai_workflow.png",16.2,"图 4-3 AI 自然语言分析流程")
    add_heading(doc,"4.4 数据库设计",2)
    add_body(doc,"数据库以 users 和 datasets 为核心。用户拥有多个数据集，数据集产生清洗记录、分析任务、AI 会话和报告；AI 会话进一步包含多条消息。关系设计既能保存完整操作过程，又便于按用户和数据集查询。")
    add_picture(doc,"diagrams/er_diagram.png",16.0,"图 4-4 数据库 E-R 图")
    db_rows=[
        ("users","用户账号、密码摘要和管理员标识","id"),
        ("datasets","文件路径、规模、字段和质量信息","id, user_id"),
        ("cleaning_records","清洗动作、参数、处理前后行数","id, dataset_id"),
        ("analysis_tasks","分析参数、结果和图表配置","id, dataset_id"),
        ("ai_conversations","数据集下的 AI 会话","id, dataset_id"),
        ("ai_messages","问题、回答和 JSON 计划","id, conversation_id"),
        ("reports","自动生成的 HTML 报告","id, dataset_id"),
    ]
    add_table(doc,["表名","主要内容","关键字段"],db_rows,[3.3,8.5,4.0],9.5)
    add_heading(doc,"4.5 部署结构",2)
    add_body(doc,"开发环境可直接使用 Flask 内置服务器；正式演示时在本机运行即可。应用通过 PyMySQL 连接 MySQL 5.7，通过文件目录保存上传数据。真实大模型属于可选外部服务，离线规则模式不需要网络。")
    add_picture(doc,"diagrams/deployment.png",15.2,"图 4-5 系统部署结构")

    # Chapter 5
    add_heading(doc,"第五章 详细设计与实现",1,page_break=True)
    add_heading(doc,"5.1 工程结构与应用工厂",2)
    add_body(doc,"项目根目录包含 app、docs、sample_data、scripts、sql、storage 和 tests。app/blueprints 存放路由，app/services 存放可测试的业务算法，templates 和 static 分别存放页面与前端资源。该结构使页面逻辑与数据处理逻辑分离。")
    add_code(doc,"""
InsightPilot/
├─ app/blueprints/        # 登录、数据集、分析、AI、报告路由
├─ app/services/          # 文件、数据、分析、AI、报告服务
├─ app/templates/         # Jinja2 页面
├─ app/static/            # 样式、脚本和本地第三方资源
├─ sample_data/           # 演示数据
├─ scripts/               # 启动、检查和报告生成脚本
├─ sql/init.sql           # MySQL 初始化
└─ tests/                 # 单元与集成测试
""")
    add_heading(doc,"5.2 用户认证模块",2)
    add_body(doc,"注册表单验证用户名长度、密码长度和两次密码一致性。密码通过 Werkzeug 的 generate_password_hash 保存摘要，登录时调用 check_password_hash 验证。所有业务页面使用 login_required，并在查询数据集时同时限定 user_id，跨用户访问将返回 404。")
    add_picture(doc,"screenshots/01_register.png",15.6,"图 5-1 用户注册页面")
    add_heading(doc,"5.3 数据集管理模块",2)
    add_body(doc,"上传路由首先验证文件是否存在、扩展名是否允许，再使用 secure_filename 和 UUID 生成存储名称。读取数据成功后，系统统计行列数、推断字段类型并生成质量报告。若数据超过最大行数限制，则删除已上传文件并提示用户。")
    add_code(doc,"""
original, stored, path, ext = save_upload(file, app.config['UPLOAD_FOLDER'])
df = read_dataframe(path)
if len(df) > app.config['MAX_ANALYSIS_ROWS']:
    Path(path).unlink(missing_ok=True)
    raise ValueError('数据行数超过限制')
dataset.columns = infer_columns(df)
dataset.quality = quality_report(df)
""")
    add_picture(doc,"screenshots/03_dataset_detail.png",12.0,"图 5-2 数据集详情和质量报告页面")
    add_heading(doc,"5.4 数据质量检测模块",2)
    for t in [
        "质量检测以 DataFrame 为输入，统计总行数、总列数、缺失单元格、重复行、唯一值数量和数值字段异常值。异常值采用四分位距 IQR 方法：当样本小于 Q1-1.5×IQR 或大于 Q3+1.5×IQR 时记为异常。该方法不要求数据服从正态分布，适合通用表格检测。",
        "健康度以 100 分为基准，根据缺失率、重复率和异常值比例扣分。评分用于快速提示，不代替专业数据质量评价。系统同时保留每个字段的详细指标，用户可以根据业务含义决定清洗方式。",
    ]: add_body(doc,t)
    add_code(doc,"""
health_score = 100
health_score -= min(45, missing_rate * 0.8)
health_score -= min(25, duplicate_rate * 0.8)
health_score -= min(20, outlier_ratio * 0.2)
health_score = max(0, round(health_score, 1))
""")
    add_picture(doc,"charts/missing_values.png",13.6,"图 5-3 示例数据缺失值检测结果")
    add_heading(doc,"5.5 数据清洗模块",2)
    add_body(doc,"清洗函数接收 operation 和 params，不直接依赖 Flask 页面，因此可独立测试。处理完成后保存为 UTF-8-SIG CSV，并更新数据集行列数、字段结构和质量报告。每次操作写入 cleaning_records，记录操作名称、参数以及处理前后行数。")
    clean_rows=[
        ("drop_duplicates","DataFrame.drop_duplicates()","删除完全重复行"),
        ("drop_missing","dropna(subset=字段)","删除指定字段缺失行"),
        ("fill_missing","均值/中位数/众数/固定值","补全缺失数据"),
        ("strip_strings","字符串 strip()","清理文本空格"),
        ("remove_outliers","IQR 上下界过滤","删除数值异常样本"),
        ("zscore","(x-均值)/标准差","标准差标准化"),
        ("minmax","(x-最小)/(最大-最小)","离差标准化"),
    ]
    add_table(doc,["操作","实现方法","作用"],clean_rows,[3.4,7.0,5.0],9.5)
    add_picture(doc,"screenshots/04_cleaning.png",15.6,"图 5-4 数据清洗页面")
    add_heading(doc,"5.6 统计分析与可视化模块",2)
    add_body(doc,"分析服务只接受 ALLOWED_ACTIONS 中的动作。每种动作都有独立参数结构和结果结构。表格型结果统一返回 columns 与 records，图表配置统一返回 title、坐标轴和 series，前端转换为 Plotly traces。分析完成后保存 analysis_tasks，便于重新查看和写入报告。")
    analysis_rows=[
        ("describe","全部字段","计数、均值、标准差、分位数等"),
        ("value_counts","一个分类字段","频数表与柱状图"),
        ("group_aggregate","分组字段+数值字段","mean/sum/count/max/min/median"),
        ("correlation","至少两个数值字段","相关矩阵与热力图"),
        ("time_trend","日期字段+数值字段","按月等频率重采样"),
        ("pivot_table","行、列、值字段","数据透视表"),
        ("linear_regression","两个数值字段","系数、截距、R²和回归线"),
        ("kmeans","至少两个数值字段","簇标签、中心和散点图"),
    ]
    add_table(doc,["分析动作","输入","输出"],analysis_rows,[3.5,5.0,7.0],9.2)
    add_picture(doc,"screenshots/05_analysis.png",15.8,"图 5-5 统计分析工作台")
    add_heading(doc,"5.7 AI 数据助手模块",2)
    add_body(doc,"AIService 首先根据数据集字段结构创建可用字段列表和数值字段列表。规则模式通过问题中的“相关、回归、聚类、趋势、按……统计、分布”等关键词生成计划；接口模式则向大模型发送字段结构、动作清单和禁止事项。模型返回内容经过 JSON 提取和动作校验。")
    add_code(doc,"""
{
  "action": "group_aggregate",
  "group_by": ["专业"],
  "metrics": [{"column": "成绩", "method": "mean"}],
  "chart_type": "bar"
}
""")
    add_body(doc,"计划通过校验后，execute_analysis 执行真实计算。随后系统把用户问题、计划和真实结果交给解释器，生成不超过指定长度的中文总结。若外部接口不可用，则自动回退到规则总结，不影响核心演示。")
    add_picture(doc,"screenshots/06_ai_assistant.png",13.2,"图 5-6 AI 数据助手页面")
    add_heading(doc,"5.8 报告生成模块",2)
    add_body(doc,"报告服务读取数据集质量信息、已保存分析任务和最近 AI 消息，生成独立 HTML。报告包含数据集概况、分析类型、结果预览、AI 记录和结论，并可作为附件下载。课程报告与系统自动报告不同：前者描述整个软件设计，后者面向某个具体数据集。")
    add_picture(doc,"screenshots/07_report.png",12.4,"图 5-7 系统生成的数据分析报告")
    add_heading(doc,"5.9 首页与数据导航",2)
    add_body(doc,"首页使用统计卡片展示数据集数量、分析任务数量、报告数量和累计行数，并列出最近数据集。导航栏统一提供首页、数据集、报告中心和退出入口，使用户可以在不同模块之间快速切换。")
    add_picture(doc,"screenshots/02_dashboard.png",15.6,"图 5-8 系统首页")

    # Chapter 6
    add_heading(doc,"第六章 系统测试",1,page_break=True)
    add_heading(doc,"6.1 测试环境与方法",2)
    env_rows=[
        ("开发语言","Python 3.10/3.11（构建环境另以 Python 3.13 完成语法和测试验证）"),
        ("Web 框架","Flask 3.0、Flask-Login、Flask-SQLAlchemy、Flask-WTF"),
        ("数据库","目标环境 MySQL 5.7；自动化测试使用内存 SQLite 隔离"),
        ("数据分析","Pandas、NumPy、Scikit-learn"),
        ("浏览器","Chrome / Edge"),
        ("测试方法","单元测试、Flask Test Client 集成测试、人工界面检查"),
    ]
    add_table(doc,["项目","说明"],env_rows,[3.3,12.4],9.8)
    add_heading(doc,"6.2 功能测试",2)
    test_rows=[
        ("T01","注册新用户","正确数据","注册成功","通过"),("T02","重复用户名","已存在用户名","提示已存在","通过"),
        ("T03","登录","正确密码","进入首页","通过"),("T04","错误密码","错误密码","提示错误","通过"),
        ("T05","上传 CSV","UTF-8 中文文件","解析并检测","通过"),("T06","非法文件","TXT 文件","拒绝上传","通过"),
        ("T07","质量检测","含缺失和重复的数据","指标正确","通过"),("T08","删除重复","执行去重","行数减少","通过"),
        ("T09","均值填补","数值缺失字段","缺失减少","通过"),("T10","IQR 处理","含极端值数据","过滤异常","通过"),
        ("T11","分组聚合","专业+成绩","返回均值表","通过"),("T12","相关分析","两个数值字段","返回矩阵","通过"),
        ("T13","线性回归","学习时长+成绩","返回 R²","通过"),("T14","同字段回归","成绩+成绩","拒绝执行","通过"),
        ("T15","K-Means","两个数值字段","返回簇标签","通过"),("T16","AI 规则模式","中文分析问题","生成 JSON 计划","通过"),
        ("T17","非法动作","非白名单 action","拒绝执行","通过"),("T18","生成报告","已有分析任务","报告保存","通过"),
        ("T19","跨用户访问","其他用户数据编号","返回 404","通过"),("T20","超大文件","超过限制","返回 413","待目标机复测"),
    ]
    add_table(doc,["编号","测试功能","输入/条件","预期结果","结果"],test_rows,[1.2,2.8,4.3,5.0,2.0],8.5)
    add_heading(doc,"6.3 自动化测试结果",2)
    add_body(doc,"项目包含分析服务单元测试和完整流程集成测试。集成测试在内存数据库中依次完成注册、登录、上传示例数据、分组聚合、AI 提问和报告生成，并验证同字段回归会被拒绝。构建时执行 python -m pytest -q，4 项测试全部通过。")
    add_code(doc,"""
$ python -m pytest -q
....                                                                     [100%]
4 passed
""")
    add_heading(doc,"6.4 数据分析结果验证",2)
    add_body(doc,"使用 student_learning.csv 作为演示数据。文件包含学号、专业、年级、学习时长、出勤率、成绩、满意度和记录日期，共 182 行，其中人为加入少量缺失值和重复行。系统检测结果与 Pandas 直接计算一致。")
    add_picture(doc,"charts/cleaning_comparison.png",12.5,"图 6-1 数据清洗前后行数对比")
    add_picture(doc,"charts/major_average_score.png",13.8,"图 6-2 不同专业平均成绩")
    add_picture(doc,"charts/study_hours_score_regression.png",13.8,"图 6-3 学习时长与成绩回归验证")
    add_picture(doc,"charts/correlation_heatmap.png",11.8,"图 6-4 相关系数矩阵验证")
    add_heading(doc,"6.5 安全与异常测试",2)
    for t in [
        "文件上传仅允许指定扩展名，存储名称由 UUID 生成，避免使用用户输入构造服务器路径。MAX_CONTENT_LENGTH 限制请求体大小，数据行数也设置上限。",
        "所有修改请求受 CSRF 保护；用户密码不以明文保存；业务查询同时检查当前用户编号。AI 计划只允许预先实现的八类动作，不允许 SQL、Python 代码、操作系统命令或任意文件路径。",
        "测试中发现并修复了两个问题：下载未清洗 Excel 时扩展名固定为 CSV；统计页面允许用户选择同一字段作为回归自变量和因变量。修复后下载名称根据实际文件后缀生成，同字段回归会返回明确错误。",
    ]: add_body(doc,t)
    bug_rows=[
        ("B01","下载文件后缀固定为 .csv","根据 active_path 后缀生成下载名","已修复"),
        ("B02","回归可选择同一字段","增加 x_col == y_col 校验","已修复"),
        ("B03","前端依赖公共 CDN","将 Bootstrap 与 Plotly.js 放入 static/vendor","已修复"),
        ("B04","无 MySQL 时不便快速预览","增加 SQLite 演示脚本","已修复"),
    ]
    add_table(doc,["编号","问题","处理方法","状态"],bug_rows,[1.2,5.2,7.4,1.8],9.2)

    # Chapter 7
    add_heading(doc,"第七章 总结与展望",1,page_break=True)
    add_heading(doc,"7.1 项目总结",2)
    for t in [
        "本项目完成了一个面向中小型表格数据的 Python Web 分析平台，将 Flask Web 开发、MySQL 数据库、Pandas 数据处理、Scikit-learn 算法、Plotly.js 可视化和大语言模型交互整合在同一工程中。系统不是传统的单表管理系统，也不是只调用聊天接口的简单套壳，而是通过“模型理解—后端校验—程序计算—模型解释”的流程保证分析结果来源可靠。",
        "通过本次实训，进一步理解了 Web 请求与响应、蓝图与应用工厂、ORM 模型、文件上传、数据质量、异常值、分组聚合、相关分析、回归和聚类等知识，也认识到大模型输出必须经过安全边界和程序校验。项目已经提供完整源代码、数据库脚本、示例数据、测试、手册、图表和报告，可用于后续适配与答辩。",
    ]: add_body(doc,t)
    add_heading(doc,"7.2 不足与展望",2)
    add_body(doc,"当前系统仍属于课程项目，主要不足包括：大数据量时仍采用内存 DataFrame；HTML 报告中的图表尚未转为静态图片；管理员统计功能较简单；真实大模型接口需要用户自行配置；缺少任务队列和细粒度审计。后续可引入异步任务、DuckDB 或 Polars、更多图表模板、Word/PDF 报告导出、分析模板分享和本地大模型。")

    add_heading(doc,"参考文献",1,page_break=True)
    refs=[
        "[1] 齐鲁师范学院信息科学与工程学院. 《Python Web 应用实训》实训结课要求及报告模板.",
        "[2] Pallets Projects. Flask Documentation.",
        "[3] Python Software Foundation. Python Documentation.",
        "[4] The pandas development team. pandas Documentation.",
        "[5] NumPy Developers. NumPy User Guide.",
        "[6] Scikit-learn Developers. Scikit-learn User Guide.",
        "[7] Oracle Corporation. MySQL 5.7 Reference Manual.",
        "[8] SQLAlchemy Authors. SQLAlchemy Documentation.",
        "[9] Plotly Technologies Inc. Plotly JavaScript Documentation.",
        "[10] Bootstrap Team. Bootstrap Documentation.",
        "[11] 《Python数据分析与可视化（第2版）-微课视频版》课程教材.",
    ]
    for ref in refs: add_body(doc,ref,no_indent=True,size=11)

    add_heading(doc,"附录",1,page_break=True)
    add_heading(doc,"附录 A 运行命令",2)
    add_code(doc,"""
# 安装依赖
pip install -r requirements.txt

# 初始化 MySQL 数据库
mysql -uroot -p < sql/init.sql
flask --app run.py init-db

# 启动系统
python run.py

# 自动化测试
python -m pytest -q
""")
    add_heading(doc,"附录 B AI 提示词与计划格式",2)
    add_body(doc,"接口模式的系统提示明确列出允许动作、字段结构和禁止事项，并要求只返回一个 JSON 对象。后端解析后再次校验，不能把提示词当作唯一安全措施。")
    add_heading(doc,"附录 C 分工建议",2)
    split_rows=[
        ("成员 1","后端框架、数据库、用户与数据集模块、部署"),
        ("成员 2","数据分析、AI 模块、前端图表、测试与报告"),
        ("共同完成","需求分析、系统联调、答辩演示和问题修改"),
    ]
    add_table(doc,["人员","建议分工"],split_rows,[3.2,12.4],10)
    add_body(doc,"若为单人完成，可将两个记录单中的第一张保留，第二张删除或标注“无组员”，并将分工描述改为本人独立完成全部模块。")

    # Core properties
    props=doc.core_properties
    props.title=TITLE
    props.subject="Python Web 应用实训报告"
    props.author="请填写姓名"
    props.keywords="Flask, Pandas, MySQL, 大语言模型, 数据可视化"
    props.comments="由 InsightPilot 项目脚本生成，提交前请填写个人信息并复核截图。"

    OUT.parent.mkdir(parents=True,exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
